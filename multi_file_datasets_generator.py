import os
import json
import pandas as pd
import logging
from typing import List, Dict, Optional
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DependencyManager:
    """依赖管理器"""
    
    @staticmethod
    def check_and_install(package_name: str, import_name: str = None) -> bool:
        """检查并尝试安装包"""
        import_name = import_name or package_name
        try:
            __import__(import_name)
            return True
        except ImportError:
            logger.warning(f"缺少依赖 {package_name}，尝试自动安装...")
            try:
                import subprocess
                subprocess.check_call(["pip", "install", package_name], 
                                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                __import__(import_name)
                logger.info(f"成功安装 {package_name}")
                return True
            except Exception as e:
                logger.error(f"安装 {package_name} 失败: {e}")
                return False

# 全局依赖检查
DEPENDENCIES = {
    'openpyxl': DependencyManager.check_and_install('openpyxl'),
    'xlrd': DependencyManager.check_and_install('xlrd'),
    'PyPDF2': DependencyManager.check_and_install('PyPDF2'),
    'python-docx': DependencyManager.check_and_install('python-docx', 'docx'),
}

class FileProcessor:
    """文件处理器基类"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.file_extension = os.path.splitext(file_path)[1].lower()
    
    def process(self) -> Optional[str]:
        """处理文件并返回内容"""
        raise NotImplementedError
    
    def validate_file(self) -> bool:
        """验证文件是否存在且可读"""
        if not os.path.exists(self.file_path):
            logger.error(f"文件不存在: {self.file_path}")
            return False
        
        if not os.access(self.file_path, os.R_OK):
            logger.error(f"文件无读取权限: {self.file_path}")
            return False
        
        return True

class ExcelProcessor(FileProcessor):
    """Excel文件处理器"""
    
    def process(self) -> Optional[str]:
        if not self.validate_file():
            return None
        
        content_parts = []
        
        # 方法1: openpyxl (推荐，支持图片检测)
        if DEPENDENCIES['openpyxl']:
            content = self._process_with_openpyxl()
            if content:
                return content
        
        # 方法2: pandas + openpyxl
        content = self._process_with_pandas('openpyxl')
        if content:
            return content
        
        # 方法3: pandas + xlrd (老版本Excel)
        if DEPENDENCIES['xlrd']:
            content = self._process_with_pandas('xlrd')
            if content:
                return content
        
        # 方法4: 最后尝试默认引擎
        content = self._process_with_pandas(None)
        if content:
            return content
        
        logger.error(f"所有Excel解析方法都失败: {self.file_name}")
        return f"[Excel文件解析失败: {self.file_name}]"
    
    def _process_with_openpyxl(self) -> Optional[str]:
        try:
            from openpyxl import load_workbook
            content_parts = []
            
            wb = load_workbook(self.file_path, data_only=True, read_only=True)
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                content_parts.append(f"=== 工作表: {sheet_name} ===")
                
                # 检测图片
                image_count = 0
                if hasattr(sheet, '_images'):
                    image_count = len(sheet._images)
                
                if image_count > 0:
                    content_parts.append(f"[包含 {image_count} 张图片/图表]")
                
                # 获取数据范围
                if sheet.max_row > 1:
                    data = []
                    max_rows = min(sheet.max_row, 100)  # 限制最大行数
                    
                    for row in sheet.iter_rows(min_row=1, max_row=max_rows, values_only=True):
                        if any(cell is not None for cell in row):
                            clean_row = [str(cell).strip() if cell is not None else '' for cell in row]
                            data.append(clean_row)
                    
                    if data:
                        # 表头
                        if len(data) > 0:
                            content_parts.append(f"列标题: {' | '.join(data[0])}")
                        
                        # 数据行
                        for i, row in enumerate(data[1:], 1):
                            if i <= 20:
                                content_parts.append(f"第{i}行: {' | '.join(row)}")
                            elif i == 21:
                                content_parts.append(f"... (总共{len(data)-1}行数据)")
                                break
                
                content_parts.append("")
            
            wb.close()
            return "\n".join(content_parts)
            
        except Exception as e:
            logger.warning(f"openpyxl处理失败: {e}")
            return None
    
    def _process_with_pandas(self, engine: Optional[str]) -> Optional[str]:
        try:
            kwargs = {'sheet_name': None}
            if engine:
                kwargs['engine'] = engine
            
            df_dict = pd.read_excel(self.file_path, **kwargs)
            content_parts = []
            
            for sheet_name, df in df_dict.items():
                content_parts.append(f"=== 工作表: {sheet_name} ===")
                if not df.empty:
                    content_parts.append(df.to_string(index=False, max_rows=50, max_cols=20))
                else:
                    content_parts.append("[空工作表]")
                content_parts.append("")
            
            return "\n".join(content_parts)
            
        except Exception as e:
            logger.warning(f"pandas({engine})处理失败: {e}")
            return None

class PDFProcessor(FileProcessor):
    """PDF文件处理器"""
    
    def process(self) -> Optional[str]:
        if not self.validate_file():
            return None
        
        if not DEPENDENCIES['PyPDF2']:
            return f"[PDF解析需要PyPDF2库: {self.file_name}]"
        
        try:
            import PyPDF2
            with open(self.file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for i, page in enumerate(reader.pages[:50]):  # 限制页数
                    try:
                        text = page.extract_text().strip()
                        if text:
                            text_parts.append(f"=== 第{i+1}页 ===\n{text}")
                    except Exception as e:
                        logger.warning(f"PDF第{i+1}页解析失败: {e}")
                        continue
                
                if text_parts:
                    return "\n\n".join(text_parts)
                else:
                    return f"[PDF文件无法提取文本: {self.file_name}]"
                    
        except Exception as e:
            logger.error(f"PDF处理失败: {e}")
            return f"[PDF文件处理失败: {self.file_name}]"

class DocxProcessor(FileProcessor):
    """Word文档处理器"""
    
    def process(self) -> Optional[str]:
        if not self.validate_file():
            return None
        
        if not DEPENDENCIES['python-docx']:
            return f"[Word解析需要python-docx库: {self.file_name}]"
        
        try:
            from docx import Document
            doc = Document(self.file_path)
            
            content_parts = []
            
            # 提取段落
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)
            
            # 提取表格
            for i, table in enumerate(doc.tables):
                content_parts.append(f"\n=== 表格{i+1} ===")
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content_parts.append(row_text)
            
            return "\n".join(content_parts) if content_parts else f"[Word文档为空: {self.file_name}]"
            
        except Exception as e:
            logger.error(f"Word文档处理失败: {e}")
            return f"[Word文档处理失败: {self.file_name}]"

class TextProcessor(FileProcessor):
    """文本文件处理器"""
    
    def process(self) -> Optional[str]:
        if not self.validate_file():
            return None
        
        for encoding in ['utf-8', 'gbk', 'gb2312', 'latin1']:
            try:
                with open(self.file_path, 'r', encoding=encoding) as f:
                    content = f.read().strip()
                return content if content else f"[空文本文件: {self.file_name}]"
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"文本文件处理失败: {e}")
                break
        
        return f"[文本文件编码识别失败: {self.file_name}]"

class CSVProcessor(FileProcessor):
    """CSV文件处理器"""
    
    def process(self) -> Optional[str]:
        if not self.validate_file():
            return None
        
        for encoding in ['utf-8', 'gbk', 'gb2312']:
            for sep in [',', '\t', ';', '|']:
                try:
                    df = pd.read_csv(self.file_path, encoding=encoding, sep=sep, nrows=1000)
                    if not df.empty:
                        return df.to_string(index=False, max_rows=50, max_cols=20)
                except Exception:
                    continue
        
        return f"[CSV文件解析失败: {self.file_name}]"

class JSONProcessor(FileProcessor):
    """JSON文件处理器"""
    
    def process(self) -> Optional[str]:
        if not self.validate_file():
            return None
        
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return json.dumps(data, ensure_ascii=False, indent=2)
        except json.JSONDecodeError as e:
            logger.error(f"JSON解析失败: {e}")
            return f"[JSON格式错误: {self.file_name}]"
        except Exception as e:
            logger.error(f"JSON文件处理失败: {e}")
            return f"[JSON文件处理失败: {self.file_name}]"

class RobustRAGDatasetGenerator:
    def __init__(self, debug=False):
        self.llm = ChatOpenAI(
            model="qwen-plus",
            api_key=os.getenv("DASHSCOPE_API_KEY"),
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
        )
        self.debug = debug
        
        if self.debug:
            logging.getLogger().setLevel(logging.DEBUG)
        else:
            # 生产模式：只显示重要信息
            logging.getLogger().setLevel(logging.INFO)
            # 禁用HTTP调试日志
            logging.getLogger("httpx").setLevel(logging.WARNING)
            logging.getLogger("httpcore").setLevel(logging.WARNING)
        
        self.processors = {
            '.txt': TextProcessor,
            '.md': TextProcessor,
            '.csv': CSVProcessor,
            '.json': JSONProcessor,
            '.xlsx': ExcelProcessor,
            '.xls': ExcelProcessor,
            '.pdf': PDFProcessor,
            '.docx': DocxProcessor,
        }
    
    def load_document_from_file(self, file_path: str) -> Optional[str]:
        """加载单个文件"""
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return None
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.processors:
            logger.warning(f"不支持的文件格式: {file_extension}")
            return None
        
        processor_class = self.processors[file_extension]
        processor = processor_class(file_path)
        
        try:
            content = processor.process()
            if content and len(content.strip()) > 10:  # 确保有实际内容
                logger.info(f"成功加载: {os.path.basename(file_path)}")
                return content
            else:
                logger.warning(f"文件内容为空或过短: {os.path.basename(file_path)}")
                return None
        except Exception as e:
            logger.error(f"文件处理异常: {file_path}, 错误: {e}")
            return None
    
    def load_documents_from_directory(self, dir_path: str) -> List[str]:
        """从目录加载所有支持的文件"""
        if not os.path.exists(dir_path):
            logger.error(f"目录不存在: {dir_path}")
            return []
        
        documents = []
        total_files = 0
        processed_files = 0
        
        for filename in os.listdir(dir_path):
            file_path = os.path.join(dir_path, filename)
            if os.path.isfile(file_path):
                file_extension = os.path.splitext(filename)[1].lower()
                if file_extension in self.processors:
                    total_files += 1
                    content = self.load_document_from_file(file_path)
                    if content:
                        documents.append(content)
                        processed_files += 1
        
        logger.info(f"目录处理完成: {processed_files}/{total_files} 个文件成功加载")
        return documents
    
    def generate_questions_from_documents(self, documents: List[str], num_questions_per_doc: int = 2) -> List[Dict]:
        """从文档生成问题"""
        if not documents:
            logger.warning("没有可用文档生成问题")
            return []
        
        question_prompt = """
基于以下文档内容，生成{num_questions}个高质量的问题。

文档内容：
{document}

要求：
1. 问题应该基于文档的核心内容
2. 避免过于简单的是非题
3. 包含不同类型的问题（事实性、分析性、应用性）

请以JSON格式返回：
{{
    "questions": [
        "问题1",
        "问题2"
    ]
}}
"""
        
        all_questions = []
        
        for i, doc in enumerate(documents):
            logger.info(f"为文档 {i+1}/{len(documents)} 生成问题...")
            
            # 截取文档内容，避免超长
            doc_content = doc[:4000] if len(doc) > 4000 else doc
            
            messages = [
                SystemMessage(content="你是一个专业的问题生成助手，擅长根据文档内容生成高质量的问题。"),
                HumanMessage(content=question_prompt.format(
                    document=doc_content,
                    num_questions=num_questions_per_doc
                ))
            ]
            
            try:
                response = self.llm.invoke(messages)
                
                # 处理不同类型的响应
                if hasattr(response, 'content'):
                    response_text = str(response.content).strip()
                else:
                    response_text = str(response).strip()
                
                logger.debug(f"LLM响应: {response_text[:200]}...")
                
                # 清理响应格式
                if response_text.startswith('```json'):
                    response_text = response_text.replace('```json', '').replace('```', '').strip()
                elif response_text.startswith('```'):
                    response_text = response_text.replace('```', '').strip()
                
                # 尝试解析JSON
                try:
                    questions_data = json.loads(response_text)
                except json.JSONDecodeError:
                    # 如果直接解析失败，尝试提取JSON部分
                    import re
                    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                    if json_match:
                        questions_data = json.loads(json_match.group())
                    else:
                        logger.error(f"无法从响应中提取JSON: {response_text}")
                        continue
                
                questions = questions_data.get('questions', [])
                
                for question in questions:
                    if isinstance(question, str) and len(question.strip()) > 5:
                        all_questions.append({
                            'question': question.strip(),
                            'source_document': doc,
                            'doc_index': i
                        })
                    elif isinstance(question, dict):
                        # 处理带type字段的问题格式
                        q_text = question.get('question', '') or question.get('content', '') or str(question)
                        if isinstance(q_text, str) and len(q_text.strip()) > 5:
                            all_questions.append({
                                'question': q_text.strip(),
                                'source_document': doc,
                                'doc_index': i
                            })
                
            except json.JSONDecodeError as e:
                logger.error(f"文档 {i+1} JSON解析失败: {e}")
                continue
            except Exception as e:
                logger.error(f"文档 {i+1} 处理失败: {e}")
                continue
        
        logger.info(f"总共生成了 {len(all_questions)} 个问题")
        return all_questions
    
    def generate_answers_and_contexts(self, questions_data: List[Dict]) -> List[Dict]:
        """生成答案和上下文"""
        if not questions_data:
            logger.warning("没有问题数据生成答案")
            return []
        
        answer_prompt = """
基于以下上下文回答问题，要求答案准确、完整。

上下文：
{context}

问题：{question}

请以JSON格式返回：
{{
    "answer": "详细的答案",
    "relevant_context": "与问题最相关的上下文片段"
}}
"""
        
        evaluation_data = []
        
        for i, item in enumerate(questions_data):
            logger.info(f"为问题 {i+1}/{len(questions_data)} 生成答案...")
            
            messages = [
                SystemMessage(content="你是一个专业的问答助手，能够基于给定上下文提供准确、详细的答案。"),
                HumanMessage(content=answer_prompt.format(
                    context=item['source_document'][:3000],  # 限制上下文长度
                    question=item['question']
                ))
            ]
            
            try:
                response = self.llm.invoke(messages)
                
                # 处理不同类型的响应
                if hasattr(response, 'content'):
                    response_text = str(response.content).strip()
                else:
                    response_text = str(response).strip()
                
                logger.debug(f"LLM答案响应: {response_text[:200]}...")
                
                if response_text.startswith('```json'):
                    response_text = response_text.replace('```json', '').replace('```', '').strip()
                elif response_text.startswith('```'):
                    response_text = response_text.replace('```', '').strip()
                
                # 尝试解析JSON
                try:
                    answer_data = json.loads(response_text)
                except json.JSONDecodeError:
                    # 如果直接解析失败，尝试提取JSON部分
                    import re
                    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                    if json_match:
                        answer_data = json.loads(json_match.group())
                    else:
                        logger.error(f"无法从答案响应中提取JSON: {response_text}")
                        continue
                
                eval_item = {
                    'question': item['question'],
                    'answer': str(answer_data.get('answer', '')),
                    'contexts': [str(answer_data.get('relevant_context', item['source_document'][:1000]))],
                    'ground_truth': str(answer_data.get('answer', ''))
                }
                
                evaluation_data.append(eval_item)
                
            except json.JSONDecodeError as e:
                logger.error(f"问题 {i+1} JSON解析失败: {e}")
                continue
            except Exception as e:
                logger.error(f"问题 {i+1} 处理失败: {e}")
                continue
        
        logger.info(f"成功生成了 {len(evaluation_data)} 个评估样本")
        return evaluation_data
    
    def save_evaluation_dataset(self, evaluation_data: List[Dict], output_path: str):
        """保存评估数据集"""
        try:
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(evaluation_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"数据集已保存到: {output_path}")
            
            # 输出统计信息
            if evaluation_data:
                logger.info(f"数据集统计: {len(evaluation_data)} 个样本")
                avg_question_len = sum(len(item['question']) for item in evaluation_data) / len(evaluation_data)
                avg_answer_len = sum(len(item['answer']) for item in evaluation_data) / len(evaluation_data)
                logger.info(f"平均问题长度: {avg_question_len:.1f} 字符")
                logger.info(f"平均答案长度: {avg_answer_len:.1f} 字符")
        
        except Exception as e:
            logger.error(f"保存数据集失败: {e}")
            raise
    
    def generate_dataset_from_file(self, input_file: str, output_file: str, num_questions: int = 5):
        """从单个文件生成数据集"""
        logger.info("=== 开始生成RAG评估数据集（单文件） ===")
        
        content = self.load_document_from_file(input_file)
        if not content:
            logger.error("文件加载失败，无法生成数据集")
            return []
        
        documents = [content]
        questions_data = self.generate_questions_from_documents(documents, num_questions)
        evaluation_data = self.generate_answers_and_contexts(questions_data)
        self.save_evaluation_dataset(evaluation_data, output_file)
        
        logger.info("=== 数据集生成完成 ===")
        return evaluation_data
    
    def generate_dataset_from_directory(self, input_dir: str, output_file: str, num_questions: int = 2):
        """从目录生成数据集"""
        logger.info("=== 开始生成RAG评估数据集（目录） ===")
        
        documents = self.load_documents_from_directory(input_dir)
        if not documents:
            logger.error("目录中没有可用文档，无法生成数据集")
            return []
        
        questions_data = self.generate_questions_from_documents(documents, num_questions)
        evaluation_data = self.generate_answers_and_contexts(questions_data)
        self.save_evaluation_dataset(evaluation_data, output_file)
        
        logger.info("=== 数据集生成完成 ===")
        return evaluation_data

if __name__ == "__main__":
    generator = RobustRAGDatasetGenerator(debug=False)  # 关闭调试模式，减少日志
    
    # 使用示例
    # 单文件处理
    # evaluation_data = generator.generate_dataset_from_file("./data/document.xlsx", "evaluation_dataset.json", num_questions=10)
    
    # 目录处理
    evaluation_data = generator.generate_dataset_from_directory("./Yili_data/", "evaluation_dataset.json", num_questions=3)
    
    print(f"生成了 {len(evaluation_data)} 个评估样本")